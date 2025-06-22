import streamlit as st
import os
from io import BytesIO
import tempfile
from typing import List, Dict, Any
import hashlib
import json

# Document processing
import PyPDF2
import docx
from openpyxl import load_workbook
import csv
from PIL import Image
import pytesseract

# Text processing and embeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain.schema import Document

# Vector database
import pinecone
from pinecone import Pinecone, ServerlessSpec

# Environment variables
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class DocumentProcessor:
    """Handles document parsing for various file types"""
    
    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            st.error(f"Error reading PDF: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(BytesIO(file_bytes))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            st.error(f"Error reading DOCX: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_txt(file_bytes: bytes) -> str:
        """Extract text from TXT file"""
        try:
            return file_bytes.decode('utf-8')
        except Exception as e:
            st.error(f"Error reading TXT: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_csv(file_bytes: bytes) -> str:
        """Extract text from CSV file"""
        try:
            content = file_bytes.decode('utf-8')
            csv_reader = csv.reader(content.splitlines())
            text = ""
            for row in csv_reader:
                text += " | ".join(row) + "\n"
            return text
        except Exception as e:
            st.error(f"Error reading CSV: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_excel(file_bytes: bytes) -> str:
        """Extract text from Excel file"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
                tmp_file.write(file_bytes)
                tmp_file.flush()
                
                workbook = load_workbook(tmp_file.name)
                text = ""
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text += f"Sheet: {sheet_name}\n"
                    for row in sheet.iter_rows(values_only=True):
                        row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                        text += row_text + "\n"
                
                os.unlink(tmp_file.name)
                return text
        except Exception as e:
            st.error(f"Error reading Excel: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_image(file_bytes: bytes) -> str:
        """Extract text from image using OCR"""
        try:
            image = Image.open(BytesIO(file_bytes))
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            st.error(f"Error reading image: {str(e)}")
            return ""
    
    def process_file(self, uploaded_file) -> str:
        """Process uploaded file and return extracted text"""
        file_extension = uploaded_file.name.split('.')[-1].lower()
        file_bytes = uploaded_file.read()
        
        if file_extension == 'pdf':
            return self.extract_text_from_pdf(file_bytes)
        elif file_extension == 'docx':
            return self.extract_text_from_docx(file_bytes)
        elif file_extension == 'txt':
            return self.extract_text_from_txt(file_bytes)
        elif file_extension == 'csv':
            return self.extract_text_from_csv(file_bytes)
        elif file_extension in ['xlsx', 'xls']:
            return self.extract_text_from_excel(file_bytes)
        elif file_extension in ['png', 'jpg', 'jpeg', 'tiff', 'bmp']:
            return self.extract_text_from_image(file_bytes)
        else:
            st.error(f"Unsupported file type: {file_extension}")
            return ""

class VectorStore:
    """Handles Pinecone vector database operations"""
    
    def __init__(self):
        self.pc = None
        self.index = None
        self.embeddings = None
        self.initialize_pinecone()
    
    def initialize_pinecone(self):
        """Initialize Pinecone client and index"""
        try:
            api_key = os.getenv("PINECONE_API_KEY")
            if not api_key:
                st.error("Pinecone API key not found. Please set PINECONE_API_KEY in your environment.")
                return
            
            self.pc = Pinecone(api_key=api_key)
            
            # Initialize OpenAI embeddings
            openai_api_key = os.getenv("OPENAI_API_KEY")
            if not openai_api_key:
                st.error("OpenAI API key not found. Please set OPENAI_API_KEY in your environment.")
                return
            
            self.embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
            
        except Exception as e:
            st.error(f"Error initializing Pinecone: {str(e)}")
    
    def create_or_get_index(self, index_name: str = "rag-documents"):
        """Create or get existing Pinecone index"""
        try:
            # Check if index exists
            existing_indexes = [index.name for index in self.pc.list_indexes()]
            
            if index_name not in existing_indexes:
                # Create new index
                self.pc.create_index(
                    name=index_name,
                    dimension=1536,  # OpenAI embeddings dimension
                    metric="cosine",
                    spec=ServerlessSpec(
                        cloud="aws",
                        region="us-east-1"
                    )
                )
                st.success(f"Created new index: {index_name}")
            
            self.index = self.pc.Index(index_name)
            return True
            
        except Exception as e:
            st.error(f"Error creating/accessing index: {str(e)}")
            return False
    
    def add_documents(self, documents: List[Document], namespace: str = "default"):
        """Add documents to Pinecone index"""
        try:
            if not self.index or not self.embeddings:
                return False
            
            # Prepare vectors for upsert
            vectors = []
            for i, doc in enumerate(documents):
                # Generate embedding
                embedding = self.embeddings.embed_query(doc.page_content)
                
                # Create unique ID
                doc_id = hashlib.md5(f"{doc.page_content[:100]}{i}".encode()).hexdigest()
                
                # Prepare metadata
                metadata = {
                    "text": doc.page_content,
                    "source": doc.metadata.get("source", "unknown"),
                    "chunk_index": i
                }
                
                vectors.append({
                    "id": doc_id,
                    "values": embedding,
                    "metadata": metadata
                })
            
            # Upsert vectors in batches
            batch_size = 100
            for i in range(0, len(vectors), batch_size):
                batch = vectors[i:i + batch_size]
                self.index.upsert(vectors=batch, namespace=namespace)
            
            return True
            
        except Exception as e:
            st.error(f"Error adding documents to vector store: {str(e)}")
            return False
    
    def similarity_search(self, query: str, k: int = 5, namespace: str = "default") -> List[str]:
        """Search for similar documents"""
        try:
            if not self.index or not self.embeddings:
                return []
            
            # Generate query embedding
            query_embedding = self.embeddings.embed_query(query)
            
            # Search in Pinecone
            results = self.index.query(
                vector=query_embedding,
                top_k=k,
                include_metadata=True,
                namespace=namespace
            )
            
            # Extract relevant text
            contexts = []
            for match in results.matches:
                if match.score > 0.7:  # Similarity threshold
                    contexts.append(match.metadata["text"])
            
            return contexts
            
        except Exception as e:
            st.error(f"Error searching vector store: {str(e)}")
            return []

class RAGSystem:
    """Main RAG system that combines retrieval and generation"""
    
    def __init__(self):
        self.vector_store = VectorStore()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", " ", ""]
        )
        self.llm = None
        self.initialize_llm()
    
    def initialize_llm(self):
        """Initialize OpenAI ChatGPT-4o"""
        try:
            openai_api_key = os.getenv("OPENAI_API_KEY")
            if openai_api_key:
                self.llm = ChatOpenAI(
                    model="gpt-4o",
                    openai_api_key=openai_api_key,
                    temperature=0.1
                )
        except Exception as e:
            st.error(f"Error initializing LLM: {str(e)}")
    
    def process_documents(self, uploaded_files, index_name: str = "rag-documents"):
        """Process uploaded documents and store in vector database"""
        if not self.vector_store.create_or_get_index(index_name):
            return False
        
        doc_processor = DocumentProcessor()
        all_documents = []
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, uploaded_file in enumerate(uploaded_files):
            status_text.text(f"Processing {uploaded_file.name}...")
            
            # Extract text from file
            text = doc_processor.process_file(uploaded_file)
            
            if text.strip():
                # Split text into chunks
                chunks = self.text_splitter.split_text(text)
                
                # Create Document objects
                for chunk in chunks:
                    doc = Document(
                        page_content=chunk,
                        metadata={"source": uploaded_file.name}
                    )
                    all_documents.append(doc)
            
            progress_bar.progress((i + 1) / len(uploaded_files))
        
        # Add documents to vector store
        status_text.text("Storing documents in vector database...")
        success = self.vector_store.add_documents(all_documents)
        
        progress_bar.empty()
        status_text.empty()
        
        if success:
            st.success(f"Successfully processed and stored {len(all_documents)} document chunks!")
        
        return success
    
    def answer_question(self, question: str, index_name: str = "rag-documents") -> str:
        """Answer question using RAG approach"""
        try:
            if not self.llm:
                return "LLM not initialized. Please check your OpenAI API key."
            
            # Retrieve relevant contexts
            contexts = self.vector_store.similarity_search(question, k=5)
            
            if not contexts:
                return "I couldn't find relevant information in the uploaded documents to answer your question."
            
            # Prepare prompt
            context_text = "\n\n".join(contexts)
            prompt = f"""
Based on the following context from the uploaded documents, please answer the question. 
If the answer is not found in the context, please say so.

Context:
{context_text}

Question: {question}

Answer:"""
            
            # Generate response
            response = self.llm.invoke(prompt)
            return response.content
            
        except Exception as e:
            return f"Error generating answer: {str(e)}"

def main():
    st.set_page_config(
        page_title="RAG Document Q&A System",
        page_icon="📚",
        layout="wide"
    )
    
    st.title("📚 RAG Document Q&A System")
    st.markdown("Upload documents and ask questions based on their content!")
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # API Keys status
        st.subheader("API Keys Status")
        openai_key = os.getenv("OPENAI_API_KEY")
        pinecone_key = os.getenv("PINECONE_API_KEY")
        
        st.write("OpenAI API:", "✅ Set" if openai_key else "❌ Not Set")
        st.write("Pinecone API:", "✅ Set" if pinecone_key else "❌ Not Set")
        
        if not openai_key or not pinecone_key:
            st.error("Please set your API keys in the .env file")
            st.code("""
# Create a .env file with:
OPENAI_API_KEY=your_openai_api_key_here
PINECONE_API_KEY=your_pinecone_api_key_here
            """)
        
        # Index name configuration
        st.subheader("Vector Database")
        index_name = st.text_input("Index Name", value="rag-documents")
    
    # Initialize RAG system
    if "rag_system" not in st.session_state:
        st.session_state.rag_system = RAGSystem()
    
    # Main interface
    tab1, tab2 = st.tabs(["📤 Upload Documents", "❓ Ask Questions"])
    
    with tab1:
        st.header("Upload Your Documents")
        
        uploaded_files = st.file_uploader(
            "Choose files to upload",
            type=['pdf', 'docx', 'txt', 'csv', 'xlsx', 'xls', 'png', 'jpg', 'jpeg'],
            accept_multiple_files=True,
            help="Supported formats: PDF, DOCX, TXT, CSV, Excel, Images"
        )
        
        if uploaded_files:
            st.write(f"Selected {len(uploaded_files)} file(s):")
            for file in uploaded_files:
                st.write(f"- {file.name} ({file.size} bytes)")
            
            if st.button("Process Documents", type="primary"):
                with st.spinner("Processing documents..."):
                    success = st.session_state.rag_system.process_documents(
                        uploaded_files, index_name
                    )
                    
                    if success:
                        st.session_state.documents_processed = True
    
    with tab2:
        st.header("Ask Questions About Your Documents")
        
        if not hasattr(st.session_state, 'documents_processed'):
            st.info("Please upload and process documents first in the 'Upload Documents' tab.")
        else:
            question = st.text_input(
                "Enter your question:",
                placeholder="What is the main topic discussed in the documents?"
            )
            
            if st.button("Get Answer", type="primary") and question:
                with st.spinner("Searching and generating answer..."):
                    answer = st.session_state.rag_system.answer_question(question, index_name)
                    
                    st.subheader("Answer:")
                    st.write(answer)
            
            # Chat history
            if "chat_history" not in st.session_state:
                st.session_state.chat_history = []
            
            if st.session_state.chat_history:
                st.subheader("Chat History")
                for i, (q, a) in enumerate(st.session_state.chat_history[-5:]):  # Show last 5
                    with st.expander(f"Q: {q[:50]}..."):
                        st.write(f"**Question:** {q}")
                        st.write(f"**Answer:** {a}")

if __name__ == "__main__":
    main()
